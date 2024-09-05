import streamlit as st
from docx import Document
from docx.shared import RGBColor
import io

def check_unbalanced_brackets(text):
    stack = []
    unbalanced = []
    missing = []
    
    for i, char in enumerate(text):
        if char == '{':
            stack.append((char, i))
        elif char == '}':
            if stack and stack[-1][0] == '{':
                stack.pop()
            else:
                unbalanced.append((char, i))
        elif char == '|':
            if not stack:
                unbalanced.append((char, i))
    
    # Check for missing closing brackets
    for char, pos in stack:
        missing.append((pos, len(text)))
    
    return sorted(unbalanced, key=lambda x: x[1]), missing

def highlight_text(text, unbalanced, missing):
    parts = []
    last_pos = 0
    
    all_positions = sorted([(pos, 'unbalanced') for _, pos in unbalanced] + 
                           [(start, 'missing_start') for start, _ in missing] + 
                           [(end, 'missing_end') for _, end in missing])
    
    for pos, type in all_positions:
        parts.append(text[last_pos:pos])
        if type == 'unbalanced':
            parts.append(f"<span style='background-color: red;'>{text[pos]}</span>")
            last_pos = pos + 1
        elif type == 'missing_start':
            parts.append("<span style='background-color: yellow;'>")
        elif type == 'missing_end':
            parts.append("</span>")
            last_pos = pos
    
    parts.append(text[last_pos:])
    return "".join(parts)

def create_word_document(text, unbalanced, missing):
    doc = Document()
    paragraph = doc.add_paragraph()
    
    last_pos = 0
    for char, pos in unbalanced:
        run = paragraph.add_run(text[last_pos:pos])
        run = paragraph.add_run(char)
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        last_pos = pos + 1
    
    for start, end in missing:
        run = paragraph.add_run(text[last_pos:start])
        run = paragraph.add_run(text[start:end])
        run.font.highlight_color = RGBColor(255, 255, 0)  # Yellow highlight
        last_pos = end
    
    run = paragraph.add_run(text[last_pos:])
    
    return doc

def main():
    st.title("Master Spin Bracket Checker")
    
    spin_text = st.text_area("Enter your master spin text here:", height=300)
    
    if st.button("Check Brackets"):
        unbalanced, missing = check_unbalanced_brackets(spin_text)
        
        if unbalanced or missing:
            st.error("Issues found in the master spin:")
            
            # Highlight unbalanced brackets in red and missing brackets in yellow
            highlighted_text = highlight_text(spin_text, unbalanced, missing)
            
            st.markdown(highlighted_text, unsafe_allow_html=True)
            
            # Create Word document
            doc = create_word_document(spin_text, unbalanced, missing)
            
            # Save document to bytes
            bio = io.BytesIO()
            doc.save(bio)
            
            # Offer download button
            st.download_button(
                label="Download Word Document",
                data=bio.getvalue(),
                file_name="master_spin_highlighted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.success("All brackets are balanced!")

if __name__ == "__main__":
    main()
